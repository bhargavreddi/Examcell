import json

from django.http.response import HttpResponse, HttpResponseNotFound, JsonResponse
from django.shortcuts import render
from django.views.decorators.csrf import csrf_protect
from docx.shared import Cm
from rest_framework import status
from django.conf import settings
from rest_framework.decorators import api_view
from rest_framework.response import Response
from Examination.serializer import StudentSerializer, DepartmentSerializer, RegulationSerializer, YearSerializer, \
    SemesterSerializer, ExaminationSerializer
from Examination.models import Student, Department, Regulation, Year, Semester, Examination, FacultyTimetable, Faculty, \
    Day, Timetable, Subject, RegisteredStudents, ExaminationType, Classroom, Arrangement, ClassroomCapacity, \
    Invigilation, InvigilationCount,UserModel
from django.conf import settings as djangoSettings

@api_view(['GET'])
def removeRegisterStudents(request,pk):
    if pk !=None:
        RegisteredStudents.objects.all().filter(examination__id=pk).delete()
        data = {
            'status': 'success'
        }
        return JsonResponse(data)
    data = {
        'status': 'failed'
    }
    return JsonResponse(data)

@api_view(['GET'])
def removeTimetable(request,pk):
    if pk !=None:
        Timetable.objects.all().filter(examination__id=pk).delete()
        data = {
            'status': 'success'
        }
        return JsonResponse(data)
    data = {
        'status': 'failed'
    }
    return JsonResponse(data)

@api_view(['GET'])
def getDepartments(request):
    if request.method == 'GET':
        snippet = Department.objects.all()
        serializer = DepartmentSerializer(snippet, many=True)
        return Response(serializer.data)
@api_view(['GET'])
def getRegulations(request):
    if request.method == 'GET':
        snippet = Regulation.objects.all()
        serializer = RegulationSerializer(snippet, many=True)
        return Response(serializer.data)

@api_view(['GET'])
def getYears(request):
    if request.method == 'GET':
        snippet = Year.objects.all()
        serializer = YearSerializer(snippet, many=True)
        return Response(serializer.data)

@api_view(['GET'])
def getSemesters(request):
    if request.method == 'GET':
        snippet = Semester.objects.all()
        serializer = SemesterSerializer(snippet, many=True)
        return Response(serializer.data)

@api_view(['GET'])
def getExaminations(request):
    if request.method == 'GET':
        snippet = Examination.objects.all()
        serializer = ExaminationSerializer(snippet, many=True)
        return Response(serializer.data)

@api_view(['POST'])
def promoteStudents(request):
    if request.method == 'POST':
        data = request.data
        dept = data.get('department')
        year = data.get('year')
        sem = data.get('semester')
        obj = Student.objects.all().filter(dept__name=dept,student_year__year=year,student_semester__semester=sem).values_list('student_id')
        if sem == 2:
            if year == 4:
                return HttpResponseNotFound()
            year = year + 1
            sem = 1
        else:
            sem = 2

        for id_tuple in obj:
            id = id_tuple[0]
            student = Student.objects.get(student_id = id)
            student.student_year = Year.objects.get(year=year)
            student.student_semester = Semester.objects.get(semester=sem)
            student.save()
        return HttpResponse()

@api_view(['POST'])
def demoteStudents(request):
    if request.method == 'POST':
        data = request.data
        dept = data.get('department')
        year = data.get('year')
        sem = data.get('semester')
        obj = Student.objects.all().filter(dept__name=dept,student_year__year=year,student_semester__semester=sem).values_list('student_id')
        if sem == 1:
            if year == 1:
                return HttpResponseNotFound()
            year = year - 1
            sem = 2
        else:
            sem = 1

        for id_tuple in obj:
            id = id_tuple[0]
            student = Student.objects.get(student_id = id)
            student.student_year = Year.objects.get(year=year)
            student.student_semester = Semester.objects.get(semester=sem)
            student.save()
        return HttpResponse()

@api_view(['POST'])
def uploadStudentList(request):
    if request.method == 'POST':
        from openpyxl import  load_workbook
        file = request.FILES.get('file')
        workbook = load_workbook(file)
        sheet_list = workbook.sheetnames
        dept_names = {'cse': 'Computer Science and Engineering','eee' : 'Electrical & Electronics Engineering','ece':'Electronics and Communication Engineering','it': 'Information Technology','mech':'Mechanical Engineering','chem':'Chemical Engineering','civil':'Civil Engineering'}
        for sheet_name in sheet_list:
            sheet = workbook[sheet_name]
            row = 2
            while True:
                if sheet.cell(row=row, column=1).value == None:
                    break
                regno = sheet.cell(row=row, column=2).value
                student_name = sheet.cell(row=row, column=3).value
                regulation = sheet.cell(row=row,column=7).value
                obj = Student()
                obj.student_id = regno
                obj.student_name = student_name
                obj.student_year = Year.objects.all().get(year='1')
                obj.student_semester = Semester.objects.all().get(semester='1')
                obj.regulation = Regulation.objects.all().get(regulation=regulation)
                obj.dept = Department.objects.all().get(name=dept_names[sheet_name])
                obj.save()
                row = row + 1
        data = {
            'status': 'success'
        }
        return JsonResponse(data)

@api_view(['POST'])
def registerStudentList(request):
    if request.method == 'POST':
        from openpyxl import load_workbook
        file = request.FILES.get('file')
        exam_name = file.__str__().split('.')
        exam_name = exam_name[0]
        RegisteredStudents.objects.all().filter(examination__examination_name=exam_name).delete()
        workbook = load_workbook(file)
        # Getting List of All Departments
        sheet_list = workbook.sheetnames
        dept_names = {'cse': 'Computer Science and Engineering', 'eee': 'Electrical & Electronics Engineering',
                          'ece': 'Electronics and Communication Engineering', 'it': 'Information Technology',
                          'mech': 'Mechanical Engineering', 'chem': 'Chemical Engineering', 'civil': 'Civil Engineering'}
        for sheet_name in sheet_list:
            sheet = workbook[sheet_name]
            row = 2
            while True:
                if sheet.cell(row=row, column=1).value == None:
                    break
                hallticket = sheet.cell(row=row, column=2).value
                year = sheet.cell(row = row,column=5).value
                semester = sheet.cell(row=row, column=6).value
                subject_code = sheet.cell(row=row,column=7).value
                regulation = sheet.cell(row=row,column=4).value
                type = sheet.cell(row = row,column=3).value
                try:
                    obj = RegisteredStudents()
                    obj.student = Student.objects.get(student_id=hallticket)
                    obj.subject = Subject.objects.get(subject_code=subject_code)
                    obj.regulation = Regulation.objects.get(regulation=regulation)
                    obj.type = ExaminationType.objects.get(type = type )
                    obj.examination = Examination.objects.get(examination_name=exam_name)
                    obj.save()
                except Exception:
                    data ={
                        'status' : 'failed'
                    }
                    return JsonResponse(data)
                row = row + 1

        data = {
            'status': 'success'
        }
        return JsonResponse(data)

@api_view(['POST'])
def uploadFacultyTimetable(request):
    from openpyxl import load_workbook
    file = request.FILES.get('file')
    # Opening the workbook
    workbook = load_workbook(file)
    sheet1 = workbook['Sheet1']
    days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday']
    FacultyTimetable.objects.all().delete()
    row = 1
    column = 1
    while True:
        while True:
            cell = sheet1.cell(row=row, column=column)
            if cell.value != None:

                faculty_name = cell.value
                if (faculty_name == 'DR B SRINIVAS'):
                    pass
                day = [None, None, None, None, None, None, None, None, None]
                for i in range(1, 7):
                    obj = FacultyTimetable()
                    try:
                        obj.faculty = Faculty.objects.get(faculty_name=faculty_name)
                    except Exception:
                        data = {
                            'status': 'failed'
                        }
                        return JsonResponse(data)
                    obj.day = Day.objects.get(day=days[i - 1])
                    day[1] = sheet1.cell(row=row + i, column=column + 1).value
                    if day[1] is None or day[1] == ' ':
                        day[1] = 'Free'
                    obj.hour1 = day[1]
                    day[2] = sheet1.cell(row=row + i, column=column + 2).value
                    if day[2] is None or day[2] == ' ':
                        if sheet1.cell(row=row + i,
                                       column=column + 1).coordinate in sheet1.merged_cells and sheet1.cell(row=row + i,
                                                                                                            column=column + 2).coordinate in sheet1.merged_cells:
                            day[2] = day[1]
                        else:
                            day[2] = 'Free'
                    obj.hour2 = day[2]
                    day[3] = sheet1.cell(row=row + i, column=column + 3).value
                    if day[3] is None or day[3] == ' ':
                        if sheet1.cell(row=row + i,
                                       column=column + 2).coordinate in sheet1.merged_cells and sheet1.cell(row=row + i,
                                                                                                            column=column + 3).coordinate in sheet1.merged_cells:
                            day[3] = day[2]
                        else:
                            day[3] = 'Free'
                    obj.hour3 = day[3]
                    day[4] = sheet1.cell(row=row + i, column=column + 4).value
                    if day[4] is None or day[4] == ' ':
                        if sheet1.cell(row=row + i,
                                       column=column + 3).coordinate in sheet1.merged_cells and sheet1.cell(row=row + i,
                                                                                                            column=column + 4).coordinate in sheet1.merged_cells:
                            day[4] = day[3]
                        else:
                            day[4] = 'Free'
                    obj.hour4 = day[4]
                    day[5] = sheet1.cell(row=row + i, column=column + 5).value
                    if day[5] is None or day[5] == ' ':
                        if sheet1.cell(row=row + i,
                                       column=column + 4).coordinate in sheet1.merged_cells and sheet1.cell(row=row + i,
                                                                                                            column=column + 5).coordinate in sheet1.merged_cells:
                            day[5] = day[4]
                        else:
                            day[5] = 'Free'
                    obj.hour5 = day[5]
                    day[6] = sheet1.cell(row=row + i, column=column + 6).value
                    if day[6] is None or day[6] == ' ':
                        if sheet1.cell(row=row + i,
                                       column=column + 5).coordinate in sheet1.merged_cells and sheet1.cell(row=row + i,
                                                                                                            column=column + 6).coordinate in sheet1.merged_cells:
                            day[6] = day[5]
                        else:
                            day[6] = 'Free'
                    obj.hour6 = day[6]
                    day[7] = sheet1.cell(row=row + i, column=column + 7).value
                    if day[7] is None or day[7] == ' ':
                        if sheet1.cell(row=row + i,
                                       column=column + 6).coordinate in sheet1.merged_cells and sheet1.cell(row=row + i,
                                                                                                            column=column + 7).coordinate in sheet1.merged_cells:
                            day[7] = day[6]
                        else:
                            day[7] = 'Free'
                    obj.hour7 = day[7]
                    day[8] = sheet1.cell(row=row + i, column=column + 8).value
                    if day[8] is None or day[8] == ' ':
                        if sheet1.cell(row=row + i,
                                       column=column + 7).coordinate in sheet1.merged_cells and sheet1.cell(row=row + i,
                                                                                                            column=column + 8).coordinate in sheet1.merged_cells:
                            day[8] = day[7]
                        else:
                            day[8] = 'Free'
                    obj.hour8 = day[8]
                    obj.save()
            else:
                row = row + 9
                column = 1
                break
            column = column + 10
        cell = sheet1.cell(row=row, column=column)
        if cell.value == None:
            break
    data = {
        'status': 'success'
    }
    return JsonResponse(data)

@api_view(['POST'])
def uploadTimetable(request):
    from openpyxl import load_workbook
    file = request.FILES.get('file')
    # Opening the workbook
    workbook = load_workbook(file)
    sheet_list = workbook.sheetnames
    for sheet_name in sheet_list:
        sheet = workbook[sheet_name]
        Timetable.objects.all().filter(examination__examination_name=sheet_name).delete()
        row = 2
        column = 1
        while True:
            cell = sheet.cell(row=row, column=column)
            if cell.value == None:
                break
            obj = Timetable()
            obj.date = sheet.cell(row = row,column=column).value
            obj.subject = Subject.objects.get(subject_code=sheet.cell(row = row,column=column + 1).value)
            obj.examination = Examination.objects.get(examination_name=sheet_name)
            obj.regulation = Regulation.objects.get(regulation=sheet.cell(row = row,column=column + 3).value)
            obj.time = sheet.cell(row=row,column=column+4).value
            row = row + 1
            obj.save()
    data = {
        'status': 'success'
    }
    return JsonResponse(data)



@api_view(['GET'])
def checkExam(request,id,yyyy,mm,dd):
    if request.method == 'GET':
        check = "{0}-{1}-{2}".format(yyyy,mm,dd)
        dates = Timetable.objects.filter(examination__id=id).values_list('date')
        for date in dates:
            if str(date[0]) == check:
                data = {
                    'status' : 'true'
                }
                return JsonResponse(data)
        data = {
            'status': 'false'
        }
        return JsonResponse(data)

@api_view(['GET'])
def isArranged(request,id,yyyy,mm,dd,hh,min):
    if request.method == 'GET':
        date = "{0}-{1}-{2}".format(yyyy,mm,dd)
        time = "{0}:{1}".format(hh, min)
        dates = Arrangement.objects.filter(examination__id=id,date=date,time=time).count()
        if dates == 0:
            data = {
                'status' : 'false'
            }
        else:
            data = {
                'status': 'true'
            }
        return JsonResponse(data)

@api_view(['GET'])
def getCapacity(request,id,yyyy,mm,dd,hh,min):
    if request.method == 'GET':
        date = "{0}-{1}-{2}".format(yyyy,mm,dd)
        time = "{0}:{1}".format(hh,min)
        subjects = Timetable.objects.filter(examination__id=id,date=date,time=time).values_list('subject')
        count = 0
        dhcount = 0
        for subject in subjects:
            temp = RegisteredStudents.objects.filter(examination__id=id,subject__subject_code=subject[0]).count()
            obj = Subject.objects.get(subject_code=subject[0])
            if obj.drawingType:
                dhcount = dhcount + temp
                mod = temp % 4
                dhcount = dhcount + 4 - mod
            else:
                count = count + temp
                mod = temp % 4
                count = count + 4 - mod

        data = {
            "capacity" : count,
            "dhcapacity" : dhcount
        }
        return JsonResponse(data)

@api_view(['GET'])
def retreiveClassrooms(request,id,yyyy,mm,dd,hh,min):
    if request.method == 'GET':
        date = "{0}-{1}-{2}".format(yyyy, mm, dd)
        time = "{0}:{1}".format(hh, min)
        classroom_list_tuple = Arrangement.objects.filter(examination__id=id, date=date, time=time).distinct().values_list('classroom')
        classroom = []
        for tuple in classroom_list_tuple:
            size = Classroom.objects.get(classroom_number=tuple[0]).capacity.size
            classroom.append({"classroom":tuple[0],'size' : size})
        data = {}
        data['classroom_list'] = classroom

        return JsonResponse(data)

@api_view(['GET'])
def retreiveSubjects(request,id,yyyy,mm,dd,hh,min):
    if request.method == 'GET':
        date = "{0}-{1}-{2}".format(yyyy, mm, dd)
        time = "{0}:{1}".format(hh, min)
        subject_list_tuple = Arrangement.objects.filter(examination__id=id, date=date, time=time).distinct().values_list('subject')
        subjects = []
        for tuple in subject_list_tuple:
            obj = Subject.objects.get(subject_code=tuple[0])
            time_obj = Timetable.objects.get(subject=obj)
            subjects.append({"subject_code":tuple[0],"subject_name" : obj.subject_name,'max_set' : time_obj.max_set,"start_set" : time_obj.set_number })
        data = {}
        data['subject_list'] = subjects

        return JsonResponse(data)

@api_view(['POST'])
def login(request):
    status = {}
    data = request.data
    try:
        obj = UserModel.objects.get(username=data['username'])
        if obj.password != data['password']:
            status['status'] = 'failed'
        else:
            status['status'] = 'success'
    except Exception:
        data['status'] = 'failed'
    return JsonResponse(status)
@api_view(['GET'])
def studentsClassroom(request,id,yyyy,mm,dd,hh,min,cs):
    if request.method == 'GET':
        date = "{0}-{1}-{2}".format(yyyy, mm, dd)
        time = "{0}:{1}".format(hh, min)
        student_list_tuple = Arrangement.objects.filter(examination__id=id, date=date, time=time,classroom__classroom_number=cs).values_list('student','set_number')
        student_list = []
        for tuple in student_list_tuple:
            student_list.append({'regno' : tuple[0],'set_number' : tuple[1]})
        data = {}
        data['student_list'] = student_list
        data['url'] = djangoSettings.STATIC_ROOT + '/files/ ' + '{0}-{1}-{2}-'.format(Examination.objects.get(id=id).examination_name, date,
                                                                                            cs)
        return JsonResponse(data)

@api_view(['GET'])
def studentsSubject(request,id,yyyy,mm,dd,hh,min,sub_code):
    if request.method == 'GET':
        date = "{0}-{1}-{2}".format(yyyy, mm, dd)
        time = "{0}:{1}".format(hh, min)
        branch = 'Computer Science and Engineering'
        student_list_tuple = Arrangement.objects.filter(examination__id=id, date=date, time=time,subject__subject_code = sub_code).values_list('student')
        student_list = []
        for tuple in student_list_tuple:
            if tuple[0] == '1':
                continue
            student_list.append({'regno':tuple[0]})
        data = {}
        data['student_list'] = student_list
        data['url'] = djangoSettings.STATIC_ROOT + '/files/ ' + '{0}-{1}-{2}-{3}-dform.docx'.format(Examination.objects.get(id=id).examination_name, date,branch,
                                                                                            sub_code)
        obj = Timetable.objects.get(subject__subject_code=sub_code)
        data['start_set'] = obj.set_number
        data['max_set'] = obj.max_set
        return JsonResponse(data)

@api_view(['GET','POST'])
def getStudentDetails(request,id,yyyy,mm,dd,hh,min,regno):
    if request.method == 'GET':
        date = "{0}-{1}-{2}".format(yyyy, mm, dd)
        time = "{0}:{1}".format(hh, min)
        student = Arrangement.objects.get(examination__id=id, date=date, time=time,student__student_id=regno)
        data  = {}
        data['attendance'] = student.attendance
        data['malpractice'] = student.malpractice
        data['blankomr'] = student.blankomr
        return JsonResponse(data)
    elif request.method == 'POST':
        data = request.data
        date = "{0}-{1}-{2}".format(yyyy, mm, dd)
        time = "{0}:{1}".format(hh, min)
        student = Arrangement.objects.get(examination__id=id, date=date, time=time, student__student_id=regno)
        student.attendance = data['attendance']
        student.malpractice = data['malpractice']
        student.blankomr = data['blankomr']
        student.save()
        data = {
            'status' : 'success'
        }
        return JsonResponse(data)


@api_view(['GET'])
def forceInvigilation(request,id,yyyy,mm,dd,hh,min,faculty):
    if request.method == 'GET':
        date = "{0}-{1}-{2}".format(yyyy, mm, dd)
        time = "{0}:{1}".format(hh, min)
        obj = Invigilation.objects.get(examination__id=id, date=date, time=time,faculty__faculty_id=faculty)
        if obj.status != 'CHANGED':
            obj.status = "FORCED"
            data = {
                'status': 'success'
            }
        else:
            data = {
                'status': 'fail'
            }
        obj.save()
        return JsonResponse(data)

@api_view(['GET'])
def changeInvigilation(request,id,yyyy,mm,dd,hh,min,faculty):
    if request.method == 'GET':
        date = "{0}-{1}-{2}".format(yyyy, mm, dd)
        time = "{0}:{1}".format(hh, min)
        obj = Invigilation.objects.get(examination__id=id, date=date, time=time, faculty__faculty_id=faculty)
        obj.status = "CHANGED"
        obj.save()
        time1 = time
        limit = 1
        import datetime
        date_temp = date
        date = date.replace("-", "")
        day = weekdays[datetime.datetime.strptime(date, "%Y%m%d").date().weekday()]
        time_array = time.split(':')
        exam_obj = Examination.objects.get(id=id)
        year = exam_obj.examination_name.split('-')
        year = year[0]
        if exam_obj.type.type == "Mid Term":
            duration = [1, 30]
        else:
            duration = [3, 0]
        min = (int(time_array[1]) + duration[1]) % 60
        hr = (int(time_array[0]) + duration[0]) + (int(time_array[1]) + duration[1]) / 60
        time2 = "{0}:{1}".format(hr, min)
        faculty_list = getInvigilationList(time1, time2,date_temp, day, year, limit, id)
        new_obj = Invigilation()
        new_obj.faculty = Faculty.objects.get(faculty_id=faculty_list[0])
        new_obj.classroom = obj.classroom
        new_obj.date = obj.date
        new_obj.time = obj.time
        new_obj.status = "PENDING"
        new_obj.examination = obj.examination
        new_obj.save()
        sendInvigilationRequest(id,faculty_list[0],date_temp,time,obj.classroom.classroom_number)
        changeInvigilationRequest(id,obj.faculty.faculty_id,date_temp,time,obj.classroom.classroom_number)
        data = {
            'faculty_id' : faculty_list[0],
            'classroom' : obj.classroom.classroom_number,
            'faculty_name' : Faculty.objects.get(faculty_id=faculty_list[0]).faculty_name
        }
        return JsonResponse(data)

@api_view(['GET'])
def acceptInvigilation(request,id,yyyy,mm,dd,hh,min,faculty):
    if request.method == 'GET':
        date = "{0}-{1}-{2}".format(yyyy, mm, dd)
        time = "{0}:{1}".format(hh, min)
        obj = Invigilation.objects.get(examination__id = id,date= date,time=time,faculty__faculty_id=faculty)
        if obj.status == 'PENDING':
            obj.status = 'ACCEPTED'
            obj.save()
            data = {
                'status' : 'success'
            }
        else:
            data = {
                'status': 'failed'
            }
        return JsonResponse(data)
@api_view(['GET'])
def rejectInvigilation(request,id,yyyy,mm,dd,hh,min,faculty):
    if request.method == 'GET':
        date = "{0}-{1}-{2}".format(yyyy, mm, dd)
        time = "{0}:{1}".format(hh, min)
        obj = Invigilation.objects.get(examination__id=id, date=date, time=time, faculty__faculty_id=faculty)
        if obj.status == 'PENDING':
            obj.status = 'REJECTED'
            obj.save()
            data = {
                'status': 'success'
            }
        else:
            data = {
                'status': 'failed'
            }
        return JsonResponse(data)

@api_view(['GET'])
def isSetAllocated(request,id,yyyy,mm,dd,hh,min):
    if request.method == 'GET':
        date = "{0}-{1}-{2}".format(yyyy, mm, dd)
        time = "{0}:{1}".format(hh, min)
        obj = Timetable.objects.filter(examination__id=id, date=date, time=time)
        count = 0
        for subject in obj:
            if subject.set_number == 0:
                count = 1
                break
        if count == 0:
            status = {
                'status' : 'true'
            }
        else:
            status = {
                'status': 'false'
            }
        return JsonResponse(status)


def assignSetNumber(exam_id,date,time):
    subjects_list = Timetable.objects.filter(examination__id=exam_id, date=date, time=time)
    for subject in subjects_list:
        dept_list = Department.objects.all()
        for dept in dept_list:
            if dept.name == 'Unknown':
                continue
            students = Arrangement.objects.filter(examination__id=exam_id, date=date, time=time,student__dept__name=dept.name)
            count = subject.set_number
            max_set = subject.max_set
            for student in students:
                student.set_number = count
                count = count + 1
                if count > max_set:
                    count = 1
                student.save()

@api_view(['POST'])
def setSubjects(request,id,yyyy,mm,dd,hh,min):
    if request.method == 'POST':
        data = request.data
        date = "{0}-{1}-{2}".format(yyyy, mm, dd)
        time = "{0}:{1}".format(hh, min)
        subject_list = data['subjects']
        for subject_dict in subject_list:
            try:
                obj = Timetable.objects.get(examination__id=id, date=date, time=time,subject__subject_code=subject_dict['subject'])
                obj.set_number = subject_dict['start_set']
                obj.max_set = subject_dict['max_set']
                obj.save()
            except Exception:
                status = {
                    'status': 'failed'
                }
                return JsonResponse(status)
        try:
            assignSetNumber(id,date,time)
            for subject_dict in subject_list:
                createDforms(id,date,time,subject_dict['subject'])
            classrooms = Arrangement.objects.filter(examination__id=id, date=date, time=time).distinct().values_list('classroom')
            for classroom in classrooms:
                createSeatingPlan(id,date,time,classroom[0])
        except Exception as e:
            status = {
                'status': 'failed'
            }
            return JsonResponse(status)
        status = {
            'status': 'success'
        }
        return JsonResponse(status)

#Invigilation Allocation
time = ["9:00","9:50","10:40","11:30","12:20","13:10","14:00","14:50","15:40"]
weekdays = ["Monday","Tuesday","Wednesday","Thursday","Friday","Saturday"]
def compare(time1,time2):
    time_array1 = time1.split(":")
    time_array2 = time2.split(":")

    if int(time_array1[0]) > int(time_array2[0]):
        return True
    elif int(time_array1[0]) == int(time_array2[0]):
        if int(time_array1[1]) == int(time_array2[1]):
            return False
        elif int(time_array1[1]) > int(time_array2[1]):
            return True
        else:
            return False
    else:
        return False
def inIndex(t):
    index = 0
    values = range(len(time))
    for i in values:
        if(compare(t,time[i]) == False):
            index = i
            break
    else:
        return 7
    return index

def outIndex(t):
    index = 0
    values = range(len(time))
    values.reverse()
    for i in values:
        if(compare(t,time[i]) == True):
            index = i
            break
    else:
        return 7
    return index



def getInvigilationList(time1,time2,date,day,year,limit,examination_id):
    fac_list = FacultyTimetable.objects.filter(day__day = day).values()
    indices = {0:'hour1',1:'hour2',2:'hour3',3:'hour4',4:'hour5',5:'hour6',6:'hour7',7:'hour8'}
    index1= outIndex(time1)
    index2 = inIndex(time2)
    count = 0
    faculty_list = []
    flag = True
    for d in fac_list:
        flag = True
        for i in range(index1,index2):
            try:
                array = d[indices[i]].split('-')
                array = array[1]
            except Exception:
                array = 'None'
            if (d[indices[i]] == 'Free' or array in year):
                pass
            else:
                flag = False
        if (flag):
            fac_obj_count = InvigilationCount.objects.get(faculty__faculty_id=d['faculty_id'])
            inv_count = Invigilation.objects.filter(faculty__faculty_id=d['faculty_id'],examination = Examination.objects.get(id = examination_id)).count()
            temp_count = Invigilation.objects.filter(faculty__faculty_id=d['faculty_id'],examination = Examination.objects.get(id = examination_id),date=date).count()
            if fac_obj_count.remaining != 0 and inv_count <2 and temp_count == 0:
                faculty_list.append(d['faculty_id'])
                count = count + 1
            if count == limit:
                break

    if count < limit:
        for d in fac_list:
            if d['faculty_id'] in faculty_list:
                pass
            else:
                inv_count = Invigilation.objects.filter(faculty__faculty_id=d['faculty_id'],
                                                        examination=Examination.objects.get(id=examination_id)).count()
                temp_count = Invigilation.objects.filter(faculty__faculty_id=d['faculty_id'],
                                                         examination=Examination.objects.get(id=examination_id),
                                                         date=date).count()
                if inv_count < 2 and temp_count == 0:
                    faculty_list.append(d['faculty_id'])
                    count = count + 1
            if count == limit:
                break
        else:
            for d in fac_list:
                temp_count = Invigilation.objects.filter(faculty__faculty_id=d['faculty_id'],
                                                         examination=Examination.objects.get(id=examination_id),
                                                         date=date).count()
                if d['faculty_id'] in faculty_list or temp_count != 0:
                    pass
                else:
                    faculty_list.append(d['faculty_id'])
                    count = count + 1
                if count == limit:
                    break
    return faculty_list

def sendInvigilationRequest(exam_id,faculty,date,time,classroom):
    pass

def changeInvigilationRequest(exam_id,faculty,date,time,classroom):
    pass

def forceInvigilationRequest(exam_id,faculty,date,time,classroom):
    pass

def allocateInvigilation(exam_id,date,time,classroom_list,drawing_hall_list):
    hall_list = classroom_list + drawing_hall_list
    limit = 0
    for hall in hall_list:
        obj = Classroom.objects.get(classroom_number=hall)
        if obj.capacity.size == 56:
            limit = limit + 2
        else:
            limit = limit + 1
    time1 = time
    import datetime
    date_temp = date
    date = date.replace("-","")
    day = weekdays[datetime.datetime.strptime(date, "%Y%m%d").date().weekday()]
    time_array = time.split(':')
    exam_obj = Examination.objects.get(id=exam_id)
    year = exam_obj.examination_name.split('-')
    year_dict = {'1': 'I', '2' : 'II' , '3' : 'III' , '4':'IV'}
    year = year_dict[year[0]]
    if exam_obj.type.type == "Mid Term":
        duration = [1,30]
    else:
        duration = [3,0]
    min = (int(time_array[1]) + duration[1]) % 60
    hr = (int(time_array[0])+ duration[0]) + (int(time_array[1]) + duration[1]) / 60
    time2 = "{0}:{1}".format(hr,min)
    faculty_list = getInvigilationList(time1, time2,date_temp, day, year, limit, exam_id)
    index = 0
    faculty_tuples = Invigilation.objects.filter(examination__id = exam_id,date = date_temp,time = time).values_list('faculty')
    Invigilation.objects.filter(examination__id=exam_id, date=date_temp, time=time).delete()
    for faculty_tuple in faculty_tuples:
        inv_count = InvigilationCount.objects.get(faculty__faculty_id=faculty_tuple[0])
        inv_count.remaining = inv_count + 1
        inv_count.assigned = inv_count - 1
        inv_count.save()
    for hall in hall_list:
        obj = Classroom.objects.get(classroom_number=hall)
        if obj.capacity.size == 56:
            for i in range(2):
                inv_obj = Invigilation()
                inv_obj.faculty = Faculty.objects.get(faculty_id=faculty_list[index])
                inv_obj.classroom = obj
                inv_obj.status = "PENDING"
                inv_obj.date = date_temp
                inv_obj.time = time
                inv_obj.examination = exam_obj
                inv_obj.save()
                inv_count = InvigilationCount.objects.get(faculty__faculty_id=faculty_list[index])
                inv_count.remaining = inv_count.remaining - 1
                inv_count.assigned = inv_count.assigned + 1
                inv_count.save()
                sendInvigilationRequest(exam_id,faculty_list[index],date_temp,time,hall)
                index = index + 1

        else:
            inv_obj = Invigilation()
            inv_obj.faculty = Faculty.objects.get(faculty_id=faculty_list[index])
            inv_obj.classroom = obj
            inv_obj.status = "PENDING"
            inv_obj.date = date_temp
            inv_obj.time = time
            inv_obj.examination = exam_obj
            inv_obj.save()
            inv_count = InvigilationCount.objects.get(faculty__faculty_id=faculty_list[index])
            inv_count.remaining = inv_count.remaining - 1
            inv_count.assigned = inv_count.assigned + 1
            inv_count.save()
            sendInvigilationRequest(exam_id,faculty_list[index], date_temp, time, hall)
            index = index + 1

#Document Generation
def createDforms(exam_id,date,time,subject_code):
    subject_obj = Timetable.objects.get(examination__id = exam_id,date = date,time=time,subject__subject_code=subject_code)
    start_set = subject_obj.set_number
    max_set = subject_obj.max_set
    examination_name = Examination.objects.get(id=exam_id).examination_name
    branch = 'Computer Science and Engineering'
    from docx import Document
    document = Document('template.docx')
    document.add_paragraph('Name of Exam: ' + examination_name)
    document.add_paragraph('Branch: '+branch)
    document.add_paragraph('Date: ' + date+'                                                                                        Time:'+time )
    document.add_paragraph('Subject Name & Subject Code: ' + Subject.objects.get(subject_code=subject_code).subject_name + ' & ' + subject_code)
    table = document.add_table(rows = 1,cols = max_set)
    cells = table.rows[0].cells
    sections = document.sections
    margin = 1
    for section in sections:
        section.top_margin = Cm(margin)
        section.bottom_margin = Cm(margin)
        section.left_margin = Cm(margin)
        section.right_margin = Cm(margin)
    for i in range(1,max_set + 1):
        cells[i-1].text = str(i)
    students = Arrangement.objects.filter(examination__id = exam_id,date = date,time = time,subject__subject_code=subject_code)
    index = start_set - 1
    cells = table.add_row().cells
    for student in students:
        if student.student.student_id == '1':
            continue
        cells[index].text = student.student.student_id
        index = index + 1
        if index == max_set:
            cells = table.add_row().cells
            index = 0
        document.save(
            djangoSettings.STATIC_ROOT + '/files/ ' + '{0}-{1}-{2}-{3}-dform.docx'.format(examination_name, date,branch,
                                                                                            subject_code))

def createHallPlan(exam_id,date,time,classroom):
    examination_name = Examination.objects.get(id=exam_id).examination_name
    branch = 'Computer Science and Engineering'
    from docx import Document
    document = Document('template.docx')
    document.add_paragraph('Name of Exam: ' + examination_name)
    document.add_paragraph('Branch: '+branch)
    document.add_paragraph('Date: ' + date+'                                                                                        Time:'+time )
    document.add_paragraph('Classroom : ' + classroom)
    classroom = Classroom.objects.get(classroom_number=classroom)
    size = classroom.capacity.size
    if size == 56:
        cols = 7
        rows = 14
    else:
        rows = 6
        cols = 8
    table = document.add_table(rows = rows,cols = cols)
    sections = document.sections
    margin = 1
    for section in sections:
        section.top_margin = Cm(margin)
        section.bottom_margin = Cm(margin)
        section.left_margin = Cm(margin)
        section.right_margin = Cm(margin)
    students = Arrangement.objects.filter(examination__id=exam_id, date=date, time=time,classroom__classroom_number=classroom)
    row_index = 0
    cols = 0
    if size == 56:
        for student in students:
            cells = table.rows[row_index].cells
            cells[cols].text = student.student.student_id
            row_index = row_index + 1
            if row_index == rows:
                row_index= 0
                cols = cols + 1
                if cols in [1,3,5]:
                    cols = cols + 1
    elif size == 36:
        for student in students:
            cells = table.rows[row_index].cells
            cells[cols].text = student.student.student_id
            row_index = row_index + 1
            if row_index == rows:
                row_index= 0
                cols = cols + 1
                if cols in [2,5]:
                    cols = cols + 1
        pass
    else:
        for student in students:
            cells = table.rows[row_index].cells
            cells[cols].text = student.student.student_id
            row_index = row_index + 1
            if row_index == rows:
                row_index= 0
                cols = cols + 1
                if cols in [2]:
                    cols = 6
        pass

        document.save(
            djangoSettings.STATIC_ROOT + '/files/ ' + '{0}-{1}-{2}-hallplan.docx'.format(examination_name, date,
                                                                                            classroom))

def createSeatingPlan(exam_id,date,time,classroom):
    examination_name = Examination.objects.get(id=exam_id).examination_name
    branch = 'Computer Science and Engineering'
    from docx import Document
    document = Document('template.docx')
    document.add_paragraph('Name of Exam: ' + examination_name)
    document.add_paragraph('Branch: '+branch)
    document.add_paragraph('Date: ' + date+'                                                                                        Time:'+time )
    document.add_paragraph('Classroom : ' + classroom)
    classroom = Classroom.objects.get(classroom_number=classroom)
    size = classroom.capacity.size
    if size == 56:
        cols = 7
        rows = 14 * 2
    else:
        rows = 6 * 2
        cols = 8
    table = document.add_table(rows = rows,cols = cols)
    sections = document.sections
    margin = 1
    for section in sections:
        section.top_margin = Cm(margin)
        section.bottom_margin = Cm(margin)
        section.left_margin = Cm(margin)
        section.right_margin = Cm(margin)
    students = Arrangement.objects.filter(examination__id=exam_id, date=date, time=time,classroom__classroom_number=classroom)
    row_index = 0
    cols = 0
    if size == 56:
        for student in students:
            cells = table.rows[row_index].cells
            cells[cols].text = student.student.student_id
            row_index = row_index + 1
            cells = table.rows[row_index].cells
            cells[cols].text = str(Arrangement.objects.get(examination__id=exam_id, date=date, time=time,student__student_id = student.student.student_id).set_number)
            row_index = row_index + 1
            if row_index == rows:
                row_index= 0
                cols = cols + 1
                if cols in [1,3,5]:
                    cols = cols + 1
    elif size == 36:
        for student in students:
            cells = table.rows[row_index].cells
            cells[cols].text = student.student.student_id
            row_index = row_index + 1
            cells = table.rows[row_index].cells
            cells[cols].text = str(Arrangement.objects.get(examination__id=exam_id, date=date, time=time,student__student_id = student.student.student_id).set_number)
            row_index = row_index + 1
            if row_index == rows:
                row_index= 0
                cols = cols + 1
                if cols in [2,5]:
                    cols = cols + 1
        pass
    else:
        for student in students:
            cells = table.rows[row_index].cells
            cells[cols].text = student.student.student_id
            row_index = row_index + 1
            cells = table.rows[row_index].cells
            cells[cols].text = str(Arrangement.objects.get(examination__id=exam_id, date=date, time=time,student__student_id = student.student.student_id).set_number)
            row_index = row_index + 1
            if row_index == rows:
                row_index= 0
                cols = cols + 1
                if cols in [2]:
                    cols = 6
        pass

    document.save(djangoSettings.STATIC_ROOT + '/files/ '+'{0}-{1}-{2}-seatingplan.docx'.format(examination_name,date,classroom))


@api_view(['POST'])
def arrangeStudents(request,id,yyyy,mm,dd,hh,min):
    if request.method == 'POST':
        date = "{0}-{1}-{2}".format(yyyy, mm, dd)
        time = "{0}:{1}".format(hh, min)
        classroom = request.data.get('classrooms')
        drawing_hall = request.data.get('drawinghalls')
        allocateInvigilation(id,date,time,classroom,drawing_hall)
        subjects_list = Timetable.objects.filter(examination__id=id, date=date, time=time).values_list('subject')
        dhsubjects =[]
        subjects = []
        for subject in subjects_list:
            obj = Subject.objects.get(subject_code=subject[0])
            if obj.drawingType:
                dhsubjects.append(subject[0])
            else:
                subjects.append(subject[0])
        Arrangement.objects.filter(examination__id=id,date=date,time=time).delete()
        hall_fill = 0
        insert_values = 0
        try:
            hall_size = Classroom.objects.get(classroom_number=drawing_hall[0]).capacity.size
            hall = drawing_hall[0]
            for dhsubject in dhsubjects:
                depts = Department.objects.values_list('name')
                for dept in depts:
                    if dept[0] == 'Unknown':
                        pass
                    list = RegisteredStudents.objects.filter(examination__id=id, subject__subject_code=dhsubject,student__dept__name=dept[0]).values_list('student')
                    for student in list:
                        obj = Arrangement()
                        obj.student = Student.objects.get(student_id=student[0])
                        obj.classtroom = Classroom.objects.get(classroom_number=hall)
                        obj.examination = Examination.objects.get(id = id)
                        obj.date = date
                        obj.time = time
                        obj.subject = Subject.objects.get(subject_code=dhsubject)
                        obj.save()
                        hall_fill =hall_fill + 1
                        insert_values = insert_values + 1
                        if hall_size == hall_fill:
                            hall_fill = 0
                            del drawing_hall[0]
                            hall_size = Classroom.objects.get(classroom_number=drawing_hall[0]).capacity.size
                            hall = drawing_hall[0]
                    if insert_values != 0:
                        temp = 4 - insert_values % 4
                        insert_values = 0
                        for i in range(temp):
                            obj = Arrangement()
                            obj.student = Student.objects.get(student_id='1')
                            obj.classroom = Classroom.objects.get(classroom_number=hall)
                            obj.examination = Examination.objects.get(id=id)
                            obj.date = date
                            obj.time = time
                            obj.subject = Subject.objects.get(subject_code=dhsubject)
                            obj.set_number = 0
                            obj.malpractice = False
                            obj.blankomr = False
                            obj.attendance = True
                            obj.save()
        except Exception:
            pass

        hall_fill = 0
        insert_values = 0
        try:
            hall_size = Classroom.objects.get(classroom_number=classroom[0]).capacity.size
            hall = classroom[0]
            for dhsubject in subjects:
                depts = Department.objects.values_list('name')
                for dept in depts:
                    if dept[0] == 'Unknown':
                        continue
                    list = RegisteredStudents.objects.filter(examination__id=id, subject__subject_code=dhsubject,
                                                             student__dept__name=dept[0]).values_list('student')
                    for student in list:
                        obj = Arrangement()
                        obj.student = Student.objects.get(student_id=student[0])
                        obj.classroom = Classroom.objects.get(classroom_number=hall)
                        obj.examination = Examination.objects.get(id=id)
                        obj.date = date
                        obj.time = time
                        obj.subject = Subject.objects.get(subject_code=dhsubject)
                        obj.set_number = 0
                        obj.malpractice = False
                        obj.blankomr = False
                        obj.attendance = True
                        obj.save()
                        hall_fill = hall_fill + 1
                        insert_values = insert_values + 1
                        if hall_size == hall_fill:
                            hall_fill = 0
                            del classroom[0]
                            hall_size = Classroom.objects.get(classroom_number=classroom[0]).capacity.size
                            hall = classroom[0]
                    if insert_values != 0:
                        temp = 4 - insert_values % 4
                        insert_values = 0
                        for i in range(temp):
                            obj = Arrangement()
                            obj.student = Student.objects.get(student_id='1')
                            obj.classroom = Classroom.objects.get(classroom_number=hall)
                            obj.examination = Examination.objects.get(id=id)
                            obj.date = date
                            obj.time = time
                            obj.subject = Subject.objects.get(subject_code=dhsubject)
                            obj.set_number = 0
                            obj.malpractice = False
                            obj.blankomr = False
                            obj.attendance = True
                            obj.save()
        except Exception as e:
            pass

        status = {
            'status' : 'success'
        }
        classrooms = Arrangement.objects.filter(examination__id=id, date=date, time=time).distinct().values_list(
            'classroom')
        for classroom in classrooms:
            createHallPlan(id, date, time, classroom[0])
        return JsonResponse(status)

@api_view(['GET'])
def getClassrooms(request,id,yyyy,mm,dd,hh,min):
    if request.method == 'GET':
        date = "{0}-{1}-{2}".format(yyyy, mm, dd)
        time = "{0}:{1}".format(hh, min)
        Arrangement.objects.filter(examination__id=id, date=date, time=time).delete()
        classroom_tuples = Arrangement.objects.filter(date=date,time=time).values_list('classroom')
        classroom_list = []
        for cls in classroom_tuples:
            classroom_list.append(cls[0])

        cls_list = []
        check = Classroom.objects.values_list('classroom_number')
        for classroom in check:
            if classroom[0] in classroom_list:
                pass
            else:
                cls_list.append(classroom[0])
        data = {}
        data['number'] = len(cls_list)
        data['classroom_list'] = []
        for classroom in cls_list:
            obj1 = Classroom.objects.get(classroom_number=classroom)
            size = obj1.capacity.size
            data['classroom_list'].append({'classroom':classroom,'size':size})
        return JsonResponse(data)

@api_view(['GET'])
def getInvigilationDetails(request,id,yyyy,mm,dd,hh,min):
    if request.method == 'GET':
        date = "{0}-{1}-{2}".format(yyyy, mm, dd)
        time = "{0}:{1}".format(hh, min)
        obj_list = Invigilation.objects.filter(examination__id=id, date=date, time=time)
        faculty_list = []
        for obj in obj_list:
            temp = {'faculty_id':obj.faculty.faculty_id,'faculty_name' : obj.faculty.faculty_name,'classroom' : obj.classroom.classroom_number,'status' : obj.status}
            faculty_list.append(temp)
        data = {}
        data['faculty_list'] = faculty_list
        return JsonResponse(data)

@api_view(['GET'])
def getExaminationsTimings(request,id):
    if request.method == 'GET':
        time = Timetable.objects.all().values_list('time').distinct()
        time_list = {}
        time_list['time_list'] = []
        for t in time:
            time_list['time_list'].append({'time':str(t[0])})

        return JsonResponse(time_list)



"""
    if request.method == 'PUT':
        snippet = Image.objects.all()
        serializer = ImageSerializer(snippet,many=True)
        return Response(serializer.data)
    elif request.method == 'POST':
        data = request.data
        student_id = data.get('student_id')
        student_name = data.get('student_name')
        image = Image.objects.get(id = id)
        user = User.objects.get(id=user_id)
        like = Likes.objects.filter(user = user ,image = image)
        if(like.count()==0):
            like = Likes(user = user,image = image)
            like.save()
        else:
            like = Likes.objects.get(user=user, image=image)
            like.delete()
        return Response()
"""