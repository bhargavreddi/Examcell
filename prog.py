from docx.shared import Cm

__author__ = 'Bhargav'
import django
import sys,os
sys.path.append("/home/bhargav/Examcell")
sys.path.append("/home/mvgrexamcell/Examcell")
os.environ["DJANGO_SETTINGS_MODULE"] = "Examcell.settings"
django.setup()
from Examination.models import Timetable, Examination, Subject, Arrangement, Classroom
from django.conf import settings as djangoSettings

def createDforms(exam_id,date,time,subject_code):
    subject_obj = Timetable.objects.get(examination__id = exam_id,date = date,time=time,subject__subject_code=subject_code)
    start_set = subject_obj.set_number
    max_set = subject_obj.max_set
    examination_name = Examination.objects.get(id=exam_id).examination_name
    branch = 'Computer Science and Engineering'
    from docx import Document
    document = Document('template.docx')
    document.add_paragraph('Name of Exam: ' + examination_name)
    document.add_paragraph('Branch: '+branch)
    document.add_paragraph('Date: ' + date+'                                                                                        Time:'+time )
    document.add_paragraph('Subject Name & Subject Code: ' + Subject.objects.get(subject_code=subject_code).subject_name + ' & ' + subject_code)
    table = document.add_table(rows = 1,cols = max_set)
    cells = table.rows[0].cells
    sections = document.sections
    margin = 1
    for section in sections:
        section.top_margin = Cm(margin)
        section.bottom_margin = Cm(margin)
        section.left_margin = Cm(margin)
        section.right_margin = Cm(margin)
    for i in range(1,max_set + 1):
        cells[i-1].text = str(i)
    students = Arrangement.objects.filter(examination__id = exam_id,date = date,time = time,subject__subject_code=subject_code)
    index = start_set - 1
    cells = table.add_row().cells
    for student in students:
        if student.student.student_id == '1':
            continue
        cells[index].text = student.student.student_id
        index = index + 1
        if index == max_set:
            cells = table.add_row().cells
            index = 0
        document.save(
            djangoSettings.STATIC_ROOT + '/files/ ' + '{0}-{1}-{2}-{3}-dform.docx'.format(examination_name, date,branch,
                                                                                            subject_code))

def createHallPlan(exam_id,date,time,classroom):
    examination_name = Examination.objects.get(id=exam_id).examination_name
    branch = 'Computer Science and Engineering'
    from docx import Document
    document = Document('template.docx')
    document.add_paragraph('Name of Exam: ' + examination_name)
    document.add_paragraph('Branch: '+branch)
    document.add_paragraph('Date: ' + date+'                                                                                        Time:'+time )
    document.add_paragraph('Classroom : ' + classroom)
    classroom = Classroom.objects.get(classroom_number=classroom)
    size = classroom.capacity.size
    if size == 56:
        cols = 7
        rows = 14
    else:
        rows = 6
        cols = 8
    table = document.add_table(rows = rows,cols = cols)
    sections = document.sections
    margin = 1
    for section in sections:
        section.top_margin = Cm(margin)
        section.bottom_margin = Cm(margin)
        section.left_margin = Cm(margin)
        section.right_margin = Cm(margin)
    students = Arrangement.objects.filter(examination__id=exam_id, date=date, time=time,classroom__classroom_number=classroom)
    row_index = 0
    cols = 0
    if size == 56:
        for student in students:
            cells = table.rows[row_index].cells
            cells[cols].text = student.student.student_id
            row_index = row_index + 1
            if row_index == rows:
                row_index= 0
                cols = cols + 1
                if cols in [1,3,5]:
                    cols = cols + 1
    elif size == 36:
        for student in students:
            cells = table.rows[row_index].cells
            cells[cols].text = student.student.student_id
            row_index = row_index + 1
            if row_index == rows:
                row_index= 0
                cols = cols + 1
                if cols in [2,5]:
                    cols = cols + 1
        pass
    else:
        for student in students:
            cells = table.rows[row_index].cells
            cells[cols].text = student.student.student_id
            row_index = row_index + 1
            if row_index == rows:
                row_index= 0
                cols = cols + 1
                if cols in [2]:
                    cols = 6
        pass

        document.save(
            djangoSettings.STATIC_ROOT + '/files/ ' + '{0}-{1}-{2}-hallplan.docx'.format(examination_name, date,
                                                                                            classroom))

def createSeatingPlan(exam_id,date,time,classroom):
    examination_name = Examination.objects.get(id=exam_id).examination_name
    branch = 'Computer Science and Engineering'
    from docx import Document
    document = Document('template.docx')
    document.add_paragraph('Name of Exam: ' + examination_name)
    document.add_paragraph('Branch: '+branch)
    document.add_paragraph('Date: ' + date+'                                                                                        Time:'+time )
    document.add_paragraph('Classroom : ' + classroom)
    classroom = Classroom.objects.get(classroom_number=classroom)
    size = classroom.capacity.size
    if size == 56:
        cols = 7
        rows = 14 * 2
    else:
        rows = 6 * 2
        cols = 8
    table = document.add_table(rows = rows,cols = cols)
    sections = document.sections
    margin = 1
    for section in sections:
        section.top_margin = Cm(margin)
        section.bottom_margin = Cm(margin)
        section.left_margin = Cm(margin)
        section.right_margin = Cm(margin)
    students = Arrangement.objects.filter(examination__id=exam_id, date=date, time=time,classroom__classroom_number=classroom)
    row_index = 0
    cols = 0
    if size == 56:
        for student in students:
            cells = table.rows[row_index].cells
            cells[cols].text = student.student.student_id
            row_index = row_index + 1
            cells = table.rows[row_index].cells
            cells[cols].text = str(Arrangement.objects.get(examination__id=exam_id, date=date, time=time,student__student_id = student.student.student_id).set_number)
            row_index = row_index + 1
            if row_index == rows:
                row_index= 0
                cols = cols + 1
                if cols in [1,3,5]:
                    cols = cols + 1
    elif size == 36:
        for student in students:
            cells = table.rows[row_index].cells
            cells[cols].text = student.student.student_id
            row_index = row_index + 1
            cells = table.rows[row_index].cells
            cells[cols].text = str(Arrangement.objects.get(examination__id=exam_id, date=date, time=time,student__student_id = student.student.student_id).set_number)
            row_index = row_index + 1
            if row_index == rows:
                row_index= 0
                cols = cols + 1
                if cols in [2,5]:
                    cols = cols + 1
        pass
    else:
        for student in students:
            cells = table.rows[row_index].cells
            cells[cols].text = student.student.student_id
            row_index = row_index + 1
            cells = table.rows[row_index].cells
            cells[cols].text = str(Arrangement.objects.get(examination__id=exam_id, date=date, time=time,student__student_id = student.student.student_id).set_number)
            row_index = row_index + 1
            if row_index == rows:
                row_index= 0
                cols = cols + 1
                if cols in [2]:
                    cols = 6
        pass

    document.save(djangoSettings.STATIC_ROOT + '/files/ '+'{0}-{1}-{2}-seatingplan.docx'.format(examination_name,date,classroom))

createDforms(1,'2017-04-03','9:30','RT42051')
createHallPlan(1,'2017-04-03','9:30','CS-10')
createSeatingPlan(1,'2017-04-03','9:30','CS-10')

